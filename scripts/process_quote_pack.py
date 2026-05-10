#!/usr/bin/env python3
import argparse
import csv
import json
import shutil
import subprocess
import tempfile
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

try:
    import openpyxl
except ImportError:
    openpyxl = None

try:
    from docx import Document
except ImportError:
    Document = None

MAX_ROWS_PER_SHEET = 120
MAX_CELLS_PER_ROW = 20
MAX_DOC_PARAGRAPHS = 120
MAX_PDF_TEXT_CHARS = 12000
MAX_TEXT_CHARS = 12000


def roundtrip_xlsx_via_libreoffice(path: Path):
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return None, "LibreOffice not installed"

    with tempfile.TemporaryDirectory(prefix="xlsx-roundtrip-") as tmp_dir:
        tmp_dir = Path(tmp_dir)
        ods_proc = subprocess.run(
            [soffice, "--headless", "--convert-to", "ods", "--outdir", str(tmp_dir), str(path)],
            capture_output=True,
            text=True,
        )
        if ods_proc.returncode != 0:
            return None, ods_proc.stderr.strip() or ods_proc.stdout.strip() or "xlsx->ods conversion failed"

        ods_path = tmp_dir / f"{path.stem}.ods"
        if not ods_path.exists():
            return None, "LibreOffice did not produce ODS output"

        xlsx_proc = subprocess.run(
            [soffice, "--headless", "--convert-to", "xlsx", "--outdir", str(tmp_dir), str(ods_path)],
            capture_output=True,
            text=True,
        )
        if xlsx_proc.returncode != 0:
            return None, xlsx_proc.stderr.strip() or xlsx_proc.stdout.strip() or "ods->xlsx conversion failed"

        repaired = tmp_dir / f"{path.stem}.xlsx"
        if not repaired.exists():
            return None, "LibreOffice did not produce repaired XLSX output"

        repaired_copy = tmp_dir / f"{path.stem}.roundtripped.xlsx"
        repaired.replace(repaired_copy)
        repaired_bytes = repaired_copy.read_bytes()

    return repaired_bytes, None


def read_manifest(path: Path):
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def pick_primary_spreadsheet(files):
    spreadsheets = [Path(f) for f in files if Path(f).suffix.lower() in {".xlsx", ".xls", ".csv"}]
    if not spreadsheets:
        return None

    def score(p: Path):
        name = p.name.lower()
        s = 0
        if "quote" in name:
            s += 8
        if "pricing" in name:
            s += 6
        if "cost" in name:
            s += 5
        if "estimate" in name:
            s += 5
        if "diamondback" in name:
            s += 3
        if "uea" in name:
            s += 2
        if "rev" in name:
            s += 1
        if p.suffix.lower() == ".xlsx":
            s += 4
        return s

    spreadsheets.sort(key=score, reverse=True)
    return spreadsheets[0]


def truncate(value, limit=500):
    text = str(value)
    return text if len(text) <= limit else text[:limit] + "…"


def read_csv_preview(path: Path):
    rows = []
    with open(path, newline="", encoding="utf-8", errors="ignore") as fh:
        reader = csv.reader(fh)
        for idx, row in enumerate(reader):
            if idx >= MAX_ROWS_PER_SHEET:
                break
            cleaned = [truncate(cell, 200) for cell in row[:MAX_CELLS_PER_ROW] if str(cell).strip()]
            if cleaned:
                rows.append(cleaned)
    return {
        "type": "csv",
        "name": path.name,
        "rows": rows,
    }


def read_xlsx_preview(path: Path):
    if openpyxl is None:
        return {"type": "xlsx", "name": path.name, "error": "openpyxl not installed"}

    workbook_note = None
    try:
        wb = openpyxl.load_workbook(path, data_only=True)
    except Exception as exc:
        repaired_bytes, repair_error = roundtrip_xlsx_via_libreoffice(path)
        if repaired_bytes is None:
            return {
                "type": "xlsx",
                "name": path.name,
                "error": f"Unable to read workbook directly ({exc}); LibreOffice round-trip fallback also failed ({repair_error})",
            }

        with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
            tmp.write(repaired_bytes)
            repaired_path = Path(tmp.name)
        try:
            wb = openpyxl.load_workbook(repaired_path, data_only=True)
            workbook_note = "Workbook required LibreOffice round-trip before preview extraction because the original XLSX was not readable by openpyxl."
        except Exception as repaired_exc:
            return {
                "type": "xlsx",
                "name": path.name,
                "error": f"Unable to read workbook directly ({exc}) or after LibreOffice round-trip ({repaired_exc})",
            }
        finally:
            repaired_path.unlink(missing_ok=True)

    sheets = []
    for ws in wb.worksheets:
        rows = []
        for idx, row in enumerate(ws.iter_rows(values_only=True)):
            if idx >= MAX_ROWS_PER_SHEET:
                break
            cleaned = [truncate(cell, 200) for cell in row[:MAX_CELLS_PER_ROW] if cell is not None and str(cell).strip()]
            if cleaned:
                rows.append(cleaned)
        sheets.append({"sheet": ws.title, "rows": rows})
    payload = {
        "type": "xlsx",
        "name": path.name,
        "sheets": sheets,
    }
    if workbook_note:
        payload["note"] = workbook_note
    return payload


def read_docx_text(path: Path):
    if Document is not None:
        doc = Document(path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
        return {
            "type": "docx",
            "name": path.name,
            "paragraphs": paragraphs[:MAX_DOC_PARAGRAPHS],
        }

    texts = []
    with zipfile.ZipFile(path) as zf:
        xml_bytes = zf.read("word/document.xml")
    root = ET.fromstring(xml_bytes)
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    for para in root.findall(".//w:p", ns):
        bits = [t.text for t in para.findall(".//w:t", ns) if t.text]
        joined = "".join(bits).strip()
        if joined:
            texts.append(joined)
    return {
        "type": "docx",
        "name": path.name,
        "paragraphs": texts[:MAX_DOC_PARAGRAPHS],
    }


def read_pdf_text(path: Path):
    text = ""
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        for page in reader.pages:
            extracted = page.extract_text() or ""
            text += extracted + "\n"
            if len(text) >= MAX_PDF_TEXT_CHARS:
                break
    except Exception as e:
        text = f"[PDF extraction unavailable: {e}]"
    return {
        "type": "pdf",
        "name": path.name,
        "text": text[:MAX_PDF_TEXT_CHARS],
    }


def read_plain_text(path: Path):
    text = path.read_text(encoding="utf-8", errors="ignore")
    return {
        "type": path.suffix.lower().lstrip("."),
        "name": path.name,
        "text": text[:MAX_TEXT_CHARS],
    }


def build_source_bundle(files):
    bundle = []
    for file_str in files:
        path = Path(file_str)
        suffix = path.suffix.lower()
        try:
            if suffix == ".csv":
                bundle.append(read_csv_preview(path))
            elif suffix == ".xlsx":
                bundle.append(read_xlsx_preview(path))
            elif suffix == ".docx":
                bundle.append(read_docx_text(path))
            elif suffix == ".pdf":
                bundle.append(read_pdf_text(path))
            elif suffix in {".txt", ".cfg", ".cfr", ".xml"}:
                bundle.append(read_plain_text(path))
            else:
                bundle.append({"type": suffix.lstrip("."), "name": path.name, "note": "preview not implemented"})
        except Exception as e:
            bundle.append({"type": suffix.lstrip("."), "name": path.name, "error": str(e)})
    return bundle


def render_docx_from_markdown(title: str, markdown_text: str, out_path: Path):
    if Document is None:
        return False, "python-docx not installed"

    doc = Document()
    doc.add_heading(title, level=1)
    for line in markdown_text.splitlines():
        if not line.strip():
            doc.add_paragraph("")
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(line)
    doc.save(out_path)
    return True, None


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--manifest", required=True)
    ap.add_argument("--model", default="gpt-5.4")
    args = ap.parse_args()

    manifest = read_manifest(Path(args.manifest))
    output_dir = Path(manifest["output_dir"])
    output_dir.mkdir(parents=True, exist_ok=True)

    files = manifest["files"]
    primary_sheet = pick_primary_spreadsheet(files)
    source_bundle = build_source_bundle(files)

    summary_data = {
        "project": manifest["batch_id"],
        "output_dir": str(output_dir),
        "quote_source": str(primary_sheet) if primary_sheet else None,
        "source_files": files,
        "source_bundle": source_bundle,
        "flags": [],
        "model": args.model,
    }

    if not primary_sheet:
        summary_data["flags"].append("No spreadsheet found in batch")

    prompt_path = output_dir / "agent_prompt.txt"
    prompt = f"""You are analysing a quote pack for UEA Diamondback. Use the supplied extracted source data to produce an accurate commercial summary.

Requirements:
- Be careful and accuracy-first, not fast.
- Determine the most likely source-of-truth pricing file.
- Extract or infer: supplier/vendor, quote/project name, revision/date, subtotal, VAT, grand total, lead time, inclusions, exclusions, assumptions, and notable risks.
- If a value is uncertain, say so explicitly rather than guessing.
- Call out discrepancies or missing source material clearly.
- Ignore any files from folders named 'old quotes - do not use' (already filtered, but keep the rule mentally).

Write two files in {output_dir}:
1. summary.md — a clean human-readable summary
2. summary_result.json — structured JSON with keys: project, vendor, quote_source, revision_or_date, subtotal, vat, total, lead_time, inclusions, exclusions, assumptions, risks, flags, source_files_used.

After writing those files, also create summary.docx from summary.md if possible.

Here is the extracted source bundle JSON:
{json.dumps(summary_data, indent=2)}
"""

    summary_data["agent_prompt_path"] = str(prompt_path)

    summary_json = output_dir / "summary_input.json"
    with open(summary_json, "w", encoding="utf-8") as f:
        json.dump(summary_data, f, indent=2)

    prompt_path.write_text(prompt, encoding="utf-8")

    print(json.dumps({
        "ok": True,
        "model": args.model,
        "output_dir": str(output_dir),
        "summary_input_json": str(summary_json),
        "agent_prompt_path": str(prompt_path),
        "primary_sheet": str(primary_sheet) if primary_sheet else None,
        "agent_prompt": prompt,
    }, indent=2))


if __name__ == "__main__":
    main()
