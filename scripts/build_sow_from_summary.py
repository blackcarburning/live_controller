#!/usr/bin/env python3
import argparse
import json
import shutil
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


DEFAULT_TEMPLATE = Path("/root/.openclaw/dropbox/inbox/SOW_Templates/CSI_SOW_ TEMPLATE 31 March 2020.docx")


def load_json(path: Path):
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def style_or_fallback(doc: Document, preferred: str, fallback: str = "Normal") -> str:
    try:
        doc.styles[preferred]
        return preferred
    except KeyError:
        return fallback


def remove_paragraph(paragraph):
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def clear_body_from_heading(doc: Document, heading_text: str = "Statement of work"):
    start_idx = None
    for idx, paragraph in enumerate(doc.paragraphs):
        style_name = paragraph.style.name if paragraph.style else ""
        if (
            paragraph.text.strip().lower() == heading_text.lower()
            and style_name.lower().startswith("heading")
        ):
            start_idx = idx
            break
    if start_idx is None:
        start_idx = 0
    for paragraph in list(doc.paragraphs[start_idx:]):
        remove_paragraph(paragraph)


def set_cover_page(doc: Document, title: str, subtitle: str):
    title_lines = [p for p in doc.paragraphs if p.style and p.style.name == "Front Page Title"]
    if len(title_lines) >= 3:
        title_lines[0].text = "STATEMENT OF WORK"
        title_lines[1].text = "FOR"
        title_lines[2].text = title
    elif len(title_lines) == 2:
        title_lines[0].text = "STATEMENT OF WORK"
        title_lines[1].text = title
    elif len(title_lines) == 1:
        title_lines[0].text = title

    # Prefer the first non-title paragraph after the cover title block as the subtitle/date line.
    insert_idx = None
    if title_lines:
        last_title = title_lines[-1]
        for idx, paragraph in enumerate(doc.paragraphs):
            if paragraph._element is last_title._element:
                insert_idx = idx + 1
                break
    if insert_idx is not None and insert_idx < len(doc.paragraphs):
        doc.paragraphs[insert_idx].text = subtitle


def add_toc_field(paragraph):
    run = paragraph.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and update field if the table of contents does not refresh automatically."

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(placeholder)
    run._r.append(fld_end)


def set_update_fields_on_open(doc: Document):
    settings = doc.settings.element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        update = OxmlElement("w:updateFields")
        update.set(qn("w:val"), "true")
        settings.append(update)
    else:
        existing.set(qn("w:val"), "true")


def add_page_break(doc: Document):
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def parse_markdown_lines(path: Path):
    return path.read_text(encoding="utf-8").splitlines()


def render_markdown_body(doc: Document, lines):
    body_style = style_or_fallback(doc, "Perpetual Body", "Normal")
    bullet_style = style_or_fallback(doc, "List Bullet", body_style)
    number_style = style_or_fallback(doc, "List Number", body_style)

    first_h1_skipped = False
    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph("")
            continue

        if stripped.startswith("# "):
            if not first_h1_skipped:
                first_h1_skipped = True
                continue
            doc.add_heading(stripped[2:].strip(), level=1)
            continue

        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=1)
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=2)
            continue

        if stripped.startswith("- "):
            doc.add_paragraph(stripped[2:].strip(), style=bullet_style)
            continue

        number_prefix = stripped.split(". ", 1)
        if len(number_prefix) == 2 and number_prefix[0].isdigit():
            doc.add_paragraph(number_prefix[1].strip(), style=number_style)
            continue

        doc.add_paragraph(stripped, style=body_style)


def refresh_docx_via_libreoffice(docx_path: Path):
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return False, "LibreOffice not installed"

    with tempfile.TemporaryDirectory(prefix="sow-refresh-") as tmp_dir:
        tmp_dir = Path(tmp_dir)
        first = subprocess.run(
            [soffice, "--headless", "--convert-to", "odt", "--outdir", str(tmp_dir), str(docx_path)],
            capture_output=True,
            text=True,
        )
        if first.returncode != 0:
            return False, first.stderr.strip() or first.stdout.strip() or "docx->odt conversion failed"

        odt_path = tmp_dir / f"{docx_path.stem}.odt"
        if not odt_path.exists():
            return False, "LibreOffice did not produce ODT output"

        second = subprocess.run(
            [soffice, "--headless", "--convert-to", "docx", "--outdir", str(tmp_dir), str(odt_path)],
            capture_output=True,
            text=True,
        )
        if second.returncode != 0:
            return False, second.stderr.strip() or second.stdout.strip() or "odt->docx conversion failed"

        refreshed = tmp_dir / f"{docx_path.stem}.docx"
        if not refreshed.exists():
            return False, "LibreOffice did not produce refreshed DOCX output"

        shutil.copy2(refreshed, docx_path)
        return True, None


def build_document(template_path: Path, summary_md: Path, summary_json: Path, out_path: Path, project_title: str | None = None):
    data = load_json(summary_json)
    project = project_title or data.get("project") or "UEA Diamondback Tape Analysis"
    revision = data.get("revision_or_date") or ""
    title = project.upper()
    subtitle = revision or "Generated from the supplied quote-pack evidence"

    doc = Document(str(template_path))
    set_cover_page(doc, title=title, subtitle=subtitle)
    clear_body_from_heading(doc)

    add_page_break(doc)
    doc.add_heading("Contents", level=1)
    add_toc_field(doc.add_paragraph())
    add_page_break(doc)

    render_markdown_body(doc, parse_markdown_lines(summary_md))
    set_update_fields_on_open(doc)
    doc.save(str(out_path))

    refreshed, refresh_note = refresh_docx_via_libreoffice(out_path)
    return {
        "ok": True,
        "output": str(out_path),
        "libreoffice_refresh": refreshed,
        "refresh_note": refresh_note,
    }


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--summary-md", required=True)
    ap.add_argument("--summary-json", required=True)
    ap.add_argument("--template", default=str(DEFAULT_TEMPLATE))
    ap.add_argument("--out", required=True)
    ap.add_argument("--project-title")
    args = ap.parse_args()

    result = build_document(
        template_path=Path(args.template),
        summary_md=Path(args.summary_md),
        summary_json=Path(args.summary_json),
        out_path=Path(args.out),
        project_title=args.project_title,
    )
    print(json.dumps(result, indent=2))


if __name__ == "__main__":
    main()
